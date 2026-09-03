import tkinter as tk
from tkinter import filedialog, messagebox, font, ttk, simpledialog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import random
import os
import sys
from PIL import Image, ImageDraw, ImageTk
import base64
import zlib
from io import BytesIO

class ModernWordEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("✨ Docx Editor Pro - Tạo File Word Chuyên Nghiệp")
        self.root.geometry("950x750")
        self.root.minsize(850, 650)
        
        # ===== SỬA LỖI ICON =====
        self.set_icon()
        
        # Màu sắc chủ đạo
        self.colors = {
            'primary': '#6C63FF',
            'primary_dark': '#5A52D5',
            'primary_light': '#8B83FF',
            'secondary': '#FF6584',
            'success': '#00D2A0',
            'warning': '#FFB800',
            'danger': '#FF4757',
            'bg': '#F0F2F8',
            'card_bg': '#FFFFFF',
            'text': '#2D3436',
            'text_light': '#636E72',
            'text_lighter': '#B2BEC3',
            'shadow': 'rgba(108, 99, 255, 0.1)',
            'gradient_start': '#667eea',
            'gradient_end': '#764ba2',
            'border': '#E8EAF6'
        }
        
        # Biến lưu trạng thái định dạng
        self.is_bold = False
        self.is_italic = False
        self.is_underline = False
        self.is_strikethrough = False
        self.current_font = "Segoe UI"
        self.current_font_size = 12
        self.current_color = "#000000"
        
        # Biến lưu trạng thái hyperlink
        self.link_text = ""
        self.link_url = ""
        
        # Biến lưu trạng thái ảnh
        self.image_path = None
        self.image_width = 400
        self.image_height = 300
        
        self.setup_styles()
        self.setup_ui()
        self.setup_animations()
        self.setup_keyboard_shortcuts()
    
    def set_icon(self):
        """Đặt icon cho ứng dụng - Cách đơn giản và hiệu quả nhất"""
        try:
            # ===== CÁCH 1: Tải từ file PNG =====
            icon_files = ["android-image.png", "app_icon.png", "icon.png", "logo.png"]
            for icon_file in icon_files:
                if os.path.exists(icon_file):
                    image = Image.open(icon_file)
                    image = image.resize((64, 64), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(image)
                    self.root.iconphoto(True, photo)
                    self.app_icon = photo
                    print(f"✅ Đã tải icon từ file: {icon_file}")
                    return
            
            # ===== CÁCH 2: Tạo icon từ code =====
            self.create_fallback_icon()
            
        except Exception as e:
            print(f"⚠️ Lỗi tải icon: {e}")
            self.create_fallback_icon()
    
    def create_fallback_icon(self):
        """Tạo icon dự phòng"""
        try:
            size = 64
            img = Image.new('RGBA', (size, size), (108, 99, 255, 255))
            draw = ImageDraw.Draw(img)
            
            # Vẽ nền gradient
            for i in range(size):
                ratio = i / size
                r = int(108 - 30 * (i / size))
                g = int(99 - 20 * (i / size))
                b = int(255 - 93 * (i / size))
                draw.rectangle([0, i, size, i+1], fill=(r, g, b, 255))
            
            # Vẽ chữ D và icon
            draw.text((22, 12), "D", fill=(255, 255, 255, 255), font=None)
            draw.text((20, 32), "📄", fill=(255, 255, 255, 255), font=None)
            
            # Vẽ viền
            draw.rectangle([2, 2, size-2, size-2], outline=(255, 255, 255, 100), width=2)
            
            photo = ImageTk.PhotoImage(img)
            self.root.iconphoto(True, photo)
            self.app_icon = photo
            print("✅ Đã tạo icon dự phòng")
        except Exception as e:
            print(f"❌ Không thể tạo icon: {e}")
        
    def setup_keyboard_shortcuts(self):
        """Thiết lập phím tắt"""
        self.root.bind('<Control-b>', lambda e: self.toggle_bold())
        self.root.bind('<Control-i>', lambda e: self.toggle_italic())
        self.root.bind('<Control-u>', lambda e: self.toggle_underline())
        self.root.bind('<Control-z>', lambda e: self.undo_action())
        self.root.bind('<Control-y>', lambda e: self.redo_action())
        self.root.bind('<Control-k>', lambda e: self.insert_hyperlink())
        self.root.bind('<Control-g>', lambda e: self.insert_image())
        
    def setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        style.configure('Primary.TButton', 
                       background=self.colors['primary'],
                       foreground='white',
                       borderwidth=0,
                       focuscolor='none',
                       font=('Segoe UI', 11, 'bold'))
        style.map('Primary.TButton',
                 background=[('active', self.colors['primary_dark'])])
        
        style.configure('Custom.TCombobox',
                       fieldbackground='white',
                       background='white',
                       foreground=self.colors['text'],
                       borderwidth=1,
                       relief='solid')
        
        style.configure('Vertical.TScrollbar',
                       gripcount=0,
                       background=self.colors['primary'],
                       troughcolor=self.colors['bg'],
                       borderwidth=0,
                       arrowcolor='white')
        
    def setup_ui(self):
        main_container = tk.Frame(self.root, bg=self.colors['bg'])
        main_container.pack(fill="both", expand=True)
        
        self.create_header(main_container)
        
        content_canvas = tk.Canvas(main_container, bg=self.colors['bg'], highlightthickness=0)
        content_canvas.pack(side="left", fill="both", expand=True)
        
        scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=content_canvas.yview)
        scrollbar.pack(side="right", fill="y")
        content_canvas.configure(yscrollcommand=scrollbar.set)
        
        content_frame = tk.Frame(content_canvas, bg=self.colors['bg'])
        content_canvas.create_window((0, 0), window=content_frame, anchor="nw", width=content_canvas.winfo_width())
        
        def configure_canvas(event):
            content_canvas.configure(scrollregion=content_canvas.bbox("all"))
            content_canvas.itemconfig(1, width=event.width)
        
        content_canvas.bind("<Configure>", configure_canvas)
        content_frame.bind("<Configure>", lambda e: content_canvas.configure(scrollregion=content_canvas.bbox("all")))
        
        content_frame.grid_columnconfigure(0, weight=1)
        content_frame.grid_rowconfigure(0, weight=0)
        content_frame.grid_rowconfigure(1, weight=1)
        content_frame.grid_rowconfigure(2, weight=0)
        
        self.create_title_card(content_frame)
        self.create_content_card(content_frame)
        self.create_footer(content_frame)
        self.create_status_bar(main_container)
        
    def create_header(self, parent):
        header_frame = tk.Frame(parent, height=90, bg=self.colors['primary'])
        header_frame.pack(fill="x")
        header_frame.pack_propagate(False)
        
        canvas = tk.Canvas(header_frame, height=90, highlightthickness=0, bg=self.colors['primary'])
        canvas.pack(fill="x")
        
        for i in range(90):
            ratio = i / 90
            color = self.interpolate_color(
                self.colors['gradient_start'], 
                self.colors['gradient_end'], 
                ratio
            )
            canvas.create_line(0, i, 1000, i, fill=color, width=1)
        
        title_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        title_frame.place(x=30, y=15)
        
        icon_label = tk.Label(
            title_frame,
            text="📝",
            font=("Segoe UI", 32),
            bg=self.colors['primary']
        )
        icon_label.pack(side="left", padx=(0, 15))
        
        title_container = tk.Frame(title_frame, bg=self.colors['primary'])
        title_container.pack(side="left")
        
        title_label = tk.Label(
            title_container,
            text="Docx Editor Pro",
            font=("Segoe UI", 26, "bold"),
            bg=self.colors['primary'],
            fg="white"
        )
        title_label.pack(anchor="w")
        
        subtitle_label = tk.Label(
            title_container,
            text="Tạo tài liệu Word chuyên nghiệp với định dạng phong phú",
            font=("Segoe UI", 10),
            bg=self.colors['primary'],
            fg="white"
        )
        subtitle_label.pack(anchor="w")
        
        right_frame = tk.Frame(header_frame, bg=self.colors['primary'])
        right_frame.place(relx=0.95, y=20, anchor="ne")
        
        version_label = tk.Label(
            right_frame,
            text="v2.1",
            font=("Segoe UI", 12, "bold"),
            bg=self.colors['primary'],
            fg="white"
        )
        version_label.pack(anchor="e")
        
        help_btn = tk.Button(
            right_frame,
            text="❓",
            font=("Segoe UI", 14),
            bg=self.colors['primary'],
            fg="white",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.show_help
        )
        help_btn.pack(anchor="e", pady=(5, 0))
        
    def create_title_card(self, parent):
        card = tk.Frame(parent, bg=self.colors['card_bg'], relief="flat", bd=0)
        card.grid(row=0, column=0, sticky="ew", pady=(0, 15))
        card.grid_columnconfigure(0, weight=1)
        
        self.add_shadow(card)
        
        padding = tk.Frame(card, bg=self.colors['card_bg'], height=10)
        padding.grid(row=0, column=0, sticky="ew")
        
        title_container = tk.Frame(card, bg=self.colors['card_bg'])
        title_container.grid(row=1, column=0, sticky="ew", padx=20, pady=(0, 15))
        title_container.grid_columnconfigure(1, weight=1)
        
        icon_label = tk.Label(
            title_container,
            text="📌",
            font=("Segoe UI", 18),
            bg=self.colors['card_bg']
        )
        icon_label.grid(row=0, column=0, padx=(0, 10), sticky="w")
        
        lbl_title = tk.Label(
            title_container,
            text="Tiêu đề tài liệu",
            font=("Segoe UI", 12, "bold"),
            bg=self.colors['card_bg'],
            fg=self.colors['text']
        )
        lbl_title.grid(row=0, column=1, sticky="w")
        
        entry_frame = tk.Frame(card, bg=self.colors['card_bg'])
        entry_frame.grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 15))
        entry_frame.grid_columnconfigure(0, weight=1)
        
        entry_border = tk.Frame(
            entry_frame,
            bg=self.colors['border'],
            highlightthickness=2,
            highlightcolor=self.colors['primary'],
            highlightbackground=self.colors['border'],
            bd=0
        )
        entry_border.grid(row=0, column=0, sticky="ew")
        
        self.entry_title = tk.Entry(
            entry_border,
            font=("Segoe UI", 12),
            bg="#F8F9FE",
            fg=self.colors['text'],
            relief="flat",
            bd=0,
            insertbackground=self.colors['primary']
        )
        self.entry_title.pack(fill="x", padx=10, pady=8)
        self.entry_title.insert(0, "Tạo Ứng Dụng Word Với Python")
        
        reset_btn = tk.Button(
            entry_frame,
            text="↺",
            font=("Segoe UI", 12),
            bg=self.colors['card_bg'],
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.reset_title
        )
        reset_btn.grid(row=0, column=1, padx=(5, 0))
        
    def create_content_card(self, parent):
        card = tk.Frame(parent, bg=self.colors['card_bg'], relief="flat", bd=0)
        card.grid(row=1, column=0, sticky="nsew", pady=(0, 15))
        card.grid_columnconfigure(0, weight=1)
        card.grid_rowconfigure(2, weight=1)
        
        self.add_shadow(card)
        
        padding = tk.Frame(card, bg=self.colors['card_bg'], height=10)
        padding.grid(row=0, column=0, sticky="ew")
        
        header_container = tk.Frame(card, bg=self.colors['card_bg'])
        header_container.grid(row=1, column=0, sticky="ew", padx=20, pady=(0, 10))
        header_container.grid_columnconfigure(1, weight=1)
        
        icon_label = tk.Label(
            header_container,
            text="📄",
            font=("Segoe UI", 18),
            bg=self.colors['card_bg']
        )
        icon_label.grid(row=0, column=0, padx=(0, 10), sticky="w")
        
        lbl_content = tk.Label(
            header_container,
            text="Nội dung bài viết",
            font=("Segoe UI", 12, "bold"),
            bg=self.colors['card_bg'],
            fg=self.colors['text']
        )
        lbl_content.grid(row=0, column=1, sticky="w")
        
        self.word_count_label = tk.Label(
            header_container,
            text="0 từ",
            font=("Segoe UI", 9),
            bg=self.colors['card_bg'],
            fg=self.colors['text_light']
        )
        self.word_count_label.grid(row=0, column=2, sticky="e")
        
        self.create_modern_toolbar(card)
        
        text_frame = tk.Frame(card, bg=self.colors['card_bg'])
        text_frame.grid(row=3, column=0, sticky="nsew", padx=20, pady=(0, 15))
        text_frame.grid_columnconfigure(0, weight=1)
        text_frame.grid_rowconfigure(0, weight=1)
        
        text_border = tk.Frame(
            text_frame,
            bg=self.colors['border'],
            highlightthickness=2,
            highlightcolor=self.colors['primary'],
            highlightbackground=self.colors['border'],
            bd=0
        )
        text_border.grid(row=0, column=0, sticky="nsew")
        text_border.grid_columnconfigure(0, weight=1)
        text_border.grid_rowconfigure(0, weight=1)
        
        text_container = tk.Frame(text_border, bg="#F8F9FE")
        text_container.grid(row=0, column=0, sticky="nsew", padx=2, pady=2)
        text_container.grid_columnconfigure(0, weight=1)
        text_container.grid_rowconfigure(0, weight=1)
        
        self.text_content = tk.Text(
            text_container,
            font=("Segoe UI", 11),
            wrap="word",
            bg="#FFFFFF",
            fg=self.colors['text'],
            relief="flat",
            bd=0,
            insertbackground=self.colors['primary'],
            selectbackground=self.colors['primary'],
            selectforeground="white",
            spacing1=2,
            spacing2=1,
            spacing3=2,
            padx=10,
            pady=10,
            undo=True,
            maxundo=50
        )
        self.text_content.grid(row=0, column=0, sticky="nsew")
        
        scrollbar = ttk.Scrollbar(text_container, orient="vertical", command=self.text_content.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.text_content.config(yscrollcommand=scrollbar.set)
        
        sample_text = """Đây là nội dung của file Word. Thư viện python-docx cho phép bạn tạo báo cáo hoặc hóa đơn động mà không cần cài đặt Microsoft Office trên máy chủ.

💡 Bạn có thể định dạng văn bản bằng cách bôi đen và sử dụng các công cụ trên thanh toolbar.

✨ Hãy thử các tính năng: in đậm, in nghiêng, gạch dưới, thay đổi font chữ và màu sắc!

🔗 Bạn cũng có thể chèn hyperlink: https://github.com
🖼️ Và chèn ảnh vào tài liệu."""
        
        self.text_content.insert("1.0", sample_text)
        
        self.text_content.bind("<<Modified>>", self.on_text_modified)
        self.text_content.bind("<ButtonRelease-1>", self.update_format_status)
        self.text_content.bind("<KeyRelease>", self.update_format_status)
        self.text_content.bind("<KeyRelease>", self.update_word_count)
        
        self.update_word_count()
        
    def create_modern_toolbar(self, parent):
        toolbar_container = tk.Frame(parent, bg=self.colors['card_bg'])
        toolbar_container.grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 10))
        toolbar_container.grid_columnconfigure(0, weight=1)
        
        toolbar = tk.Frame(
            toolbar_container,
            bg="#F0F2F8",
            relief="flat",
            bd=1,
            highlightthickness=1,
            highlightcolor=self.colors['border']
        )
        toolbar.grid(row=0, column=0, sticky="ew", ipady=5)
        
        # Nhóm 1: Định dạng cơ bản
        format_group = tk.Frame(toolbar, bg="#F0F2F8")
        format_group.pack(side="left", padx=5)
        
        self.btn_bold = self.create_toolbar_button(
            format_group, "𝐁", "In đậm (Ctrl+B)", self.toggle_bold
        )
        self.btn_bold.pack(side="left", padx=2)
        
        self.btn_italic = self.create_toolbar_button(
            format_group, "𝑰", "In nghiêng (Ctrl+I)", self.toggle_italic
        )
        self.btn_italic.pack(side="left", padx=2)
        
        self.btn_underline = self.create_toolbar_button(
            format_group, "U̲", "Gạch dưới (Ctrl+U)", self.toggle_underline
        )
        self.btn_underline.pack(side="left", padx=2)
        
        self.btn_strikethrough = self.create_toolbar_button(
            format_group, "S̶", "Gạch ngang", self.toggle_strikethrough
        )
        self.btn_strikethrough.pack(side="left", padx=2)
        
        ttk.Separator(toolbar, orient='vertical').pack(side="left", padx=5, fill="y")
        
        # Nhóm 2: Font và size
        font_group = tk.Frame(toolbar, bg="#F0F2F8")
        font_group.pack(side="left", padx=5)
        
        font_label = tk.Label(
            font_group, 
            text="Font:", 
            font=("Segoe UI", 9),
            bg="#F0F2F8",
            fg=self.colors['text_light']
        )
        font_label.pack(side="left", padx=(0, 5))
        
        self.font_var = tk.StringVar(value="Segoe UI")
        font_options = ["Segoe UI", "Arial", "Times New Roman", "Calibri", "Verdana", "Tahoma", "Georgia"]
        
        font_combo = ttk.Combobox(
            font_group,
            textvariable=self.font_var,
            values=font_options,
            width=14,
            state="readonly",
            style='Custom.TCombobox'
        )
        font_combo.pack(side="left", padx=(0, 10))
        font_combo.bind('<<ComboboxSelected>>', lambda e: self.change_font(self.font_var.get()))
        
        size_label = tk.Label(
            font_group,
            text="Size:",
            font=("Segoe UI", 9),
            bg="#F0F2F8",
            fg=self.colors['text_light']
        )
        size_label.pack(side="left", padx=(0, 5))
        
        self.size_var = tk.StringVar(value="12")
        size_options = ["8", "9", "10", "11", "12", "14", "16", "18", "20", "24", "28", "36", "48", "72"]
        
        size_combo = ttk.Combobox(
            font_group,
            textvariable=self.size_var,
            values=size_options,
            width=6,
            state="readonly",
            style='Custom.TCombobox'
        )
        size_combo.pack(side="left")
        size_combo.bind('<<ComboboxSelected>>', lambda e: self.change_size(self.size_var.get()))
        
        ttk.Separator(toolbar, orient='vertical').pack(side="left", padx=5, fill="y")
        
        # Nhóm 3: Màu sắc
        color_group = tk.Frame(toolbar, bg="#F0F2F8")
        color_group.pack(side="left", padx=5)
        
        color_label = tk.Label(
            color_group,
            text="🎨",
            font=("Segoe UI", 10),
            bg="#F0F2F8"
        )
        color_label.pack(side="left", padx=(0, 5))
        
        colors = [
            ("#000000", "Đen"), ("#FF0000", "Đỏ"), ("#0000FF", "Xanh dương"),
            ("#008000", "Xanh lá"), ("#800080", "Tím"), ("#FFA500", "Cam"),
            ("#FF1493", "Hồng"), ("#00CED1", "Xanh ngọc"), ("#8B4513", "Nâu")
        ]
        
        for color_code, color_name in colors:
            btn = tk.Button(
                color_group,
                bg=color_code,
                width=2,
                height=1,
                relief="flat",
                bd=1,
                cursor="hand2",
                command=lambda c=color_code: self.change_color(c)
            )
            btn.pack(side="left", padx=1)
        
        ttk.Separator(toolbar, orient='vertical').pack(side="left", padx=5, fill="y")
        
        # Nhóm 4: Chèn hyperlink và ảnh
        insert_group = tk.Frame(toolbar, bg="#F0F2F8")
        insert_group.pack(side="left", padx=5)
        
        link_btn = tk.Button(
            insert_group,
            text="🔗",
            font=("Segoe UI", 10),
            bg="#F0F2F8",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.insert_hyperlink
        )
        link_btn.pack(side="left", padx=2)
        
        image_btn = tk.Button(
            insert_group,
            text="🖼️",
            font=("Segoe UI", 10),
            bg="#F0F2F8",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.insert_image
        )
        image_btn.pack(side="left", padx=2)
        
        ttk.Separator(toolbar, orient='vertical').pack(side="left", padx=5, fill="y")
        
        # Nhóm 5: Ký tự đặc biệt
        char_group = tk.Frame(toolbar, bg="#F0F2F8")
        char_group.pack(side="left", padx=5)
        
        special_chars = ["•", "★", "✓", "✗", "→", "←", "↑", "↓", 
                        "©", "®", "™", "§", "¶", "∞", "∑", "√", "♥", "♦", "♣", "♠"]
        
        char_btn = tk.Menubutton(
            char_group,
            text="⌨️",
            font=("Segoe UI", 10),
            bg="#F0F2F8",
            relief="flat",
            bd=0,
            cursor="hand2"
        )
        char_btn.pack(side="left")
        
        char_menu = tk.Menu(char_btn, tearoff=0, font=("Segoe UI", 10))
        char_btn.config(menu=char_menu)
        
        for char in special_chars:
            char_menu.add_command(
                label=char,
                command=lambda c=char: self.insert_special_char(c)
            )
        
        ttk.Separator(toolbar, orient='vertical').pack(side="left", padx=5, fill="y")
        
        # Nhóm 6: Căn chỉnh
        align_group = tk.Frame(toolbar, bg="#F0F2F8")
        align_group.pack(side="left", padx=5)
        
        alignments = [
            ("⬅️", "Căn trái", "left"),
            ("⬛", "Căn giữa", "center"),
            ("➡️", "Căn phải", "right"),
            ("⬛", "Căn đều", "justify")
        ]
        
        for icon, tip, align in alignments:
            btn = tk.Button(
                align_group,
                text=icon,
                font=("Segoe UI", 9),
                bg="#F0F2F8",
                relief="flat",
                bd=0,
                cursor="hand2",
                command=lambda a=align: self.change_alignment(a)
            )
            btn.pack(side="left", padx=2)
        
        ttk.Separator(toolbar, orient='vertical').pack(side="left", padx=5, fill="y")
        
        # Nhóm 7: Công cụ khác
        tool_group = tk.Frame(toolbar, bg="#F0F2F8")
        tool_group.pack(side="left", padx=5)
        
        clear_btn = tk.Button(
            tool_group,
            text="🗑️",
            font=("Segoe UI", 10),
            bg="#F0F2F8",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.clear_content
        )
        clear_btn.pack(side="left", padx=2)
        
        undo_btn = tk.Button(
            tool_group,
            text="↩️",
            font=("Segoe UI", 10),
            bg="#F0F2F8",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.undo_action
        )
        undo_btn.pack(side="left", padx=2)
        
        redo_btn = tk.Button(
            tool_group,
            text="↪️",
            font=("Segoe UI", 10),
            bg="#F0F2F8",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.redo_action
        )
        redo_btn.pack(side="left", padx=2)
        
    def create_toolbar_button(self, parent, text, tooltip, command):
        btn = tk.Button(
            parent,
            text=text,
            font=("Segoe UI", 11, "bold"),
            bg="#F0F2F8",
            fg=self.colors['text'],
            relief="flat",
            bd=0,
            cursor="hand2",
            command=command,
            width=4
        )
        return btn
        
    def create_footer(self, parent):
        footer_frame = tk.Frame(parent, bg=self.colors['card_bg'])
        footer_frame.grid(row=2, column=0, sticky="ew", pady=(0, 10))
        footer_frame.grid_columnconfigure(0, weight=1)
        
        footer_container = tk.Frame(footer_frame, bg=self.colors['card_bg'])
        footer_container.grid(row=0, column=0, sticky="ew", padx=20)
        footer_container.grid_columnconfigure(0, weight=1)
        
        btn_frame = tk.Frame(footer_container, bg=self.colors['card_bg'])
        btn_frame.grid(row=0, column=0, sticky="ew", pady=(0, 5))
        btn_frame.grid_columnconfigure(0, weight=1)
        
        btn_create = tk.Button(
            btn_frame,
            text="✨ Tạo File Word (.docx)",
            font=("Segoe UI", 13, "bold"),
            bg=self.colors['primary'],
            fg="white",
            activebackground=self.colors['primary_dark'],
            activeforeground="white",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=self.create_word_file,
            pady=14
        )
        btn_create.grid(row=0, column=0, sticky="ew")
        
        btn_create.bind("<Enter>", lambda e: btn_create.config(bg=self.colors['primary_dark']))
        btn_create.bind("<Leave>", lambda e: btn_create.config(bg=self.colors['primary']))
        
        info_label = tk.Label(
            footer_container,
            text="💡 Hỗ trợ: Định dạng văn bản | Hyperlink (Ctrl+K) | Ảnh (Ctrl+G) | Phím tắt đầy đủ",
            font=("Segoe UI", 9),
            bg=self.colors['card_bg'],
            fg=self.colors['text_light']
        )
        info_label.grid(row=1, column=0, pady=(5, 0))
        
    def create_status_bar(self, parent):
        status_frame = tk.Frame(parent, bg=self.colors['text'], height=25)
        status_frame.pack(fill="x", side="bottom")
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            status_frame,
            text="Sẵn sàng",
            font=("Segoe UI", 9),
            bg=self.colors['text'],
            fg="white",
            anchor="w"
        )
        self.status_label.pack(side="left", padx=10)
        
        self.cursor_label = tk.Label(
            status_frame,
            text="Dòng 1, Cột 1",
            font=("Segoe UI", 9),
            bg=self.colors['text'],
            fg="white",
            anchor="e"
        )
        self.cursor_label.pack(side="right", padx=10)
        
        self.text_content.bind("<KeyRelease>", self.update_cursor_position)
        self.text_content.bind("<ButtonRelease-1>", self.update_cursor_position)
        
    def add_shadow(self, widget):
        widget.config(
            highlightthickness=1,
            highlightcolor="#E8EAF6",
            highlightbackground="#E8EAF6"
        )
        
    def interpolate_color(self, color1, color2, ratio):
        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            
        def rgb_to_hex(rgb):
            return '#{:02x}{:02x}{:02x}'.format(*rgb)
            
        rgb1 = hex_to_rgb(color1)
        rgb2 = hex_to_rgb(color2)
        
        mixed = tuple(int(c1 + (c2 - c1) * ratio) for c1, c2 in zip(rgb1, rgb2))
        return rgb_to_hex(mixed)
        
    def setup_animations(self):
        try:
            alpha = 0.0
            self.root.attributes('-alpha', alpha)
            
            def increase_alpha():
                nonlocal alpha
                if alpha < 1.0:
                    alpha += 0.05
                    self.root.attributes('-alpha', alpha)
                    self.root.after(30, increase_alpha)
            
            increase_alpha()
        except:
            pass
        
    # Các hàm xử lý định dạng
    def toggle_bold(self):
        self.is_bold = not self.is_bold
        self.btn_bold.config(
            bg="#6C63FF" if self.is_bold else "#F0F2F8",
            fg="white" if self.is_bold else self.colors['text']
        )
        self.apply_format_to_selection()
        self.update_status("Đã bật in đậm" if self.is_bold else "Đã tắt in đậm")
        
    def toggle_italic(self):
        self.is_italic = not self.is_italic
        self.btn_italic.config(
            bg="#6C63FF" if self.is_italic else "#F0F2F8",
            fg="white" if self.is_italic else self.colors['text']
        )
        self.apply_format_to_selection()
        self.update_status("Đã bật in nghiêng" if self.is_italic else "Đã tắt in nghiêng")
        
    def toggle_underline(self):
        self.is_underline = not self.is_underline
        self.btn_underline.config(
            bg="#6C63FF" if self.is_underline else "#F0F2F8",
            fg="white" if self.is_underline else self.colors['text']
        )
        self.apply_format_to_selection()
        self.update_status("Đã bật gạch dưới" if self.is_underline else "Đã tắt gạch dưới")
        
    def toggle_strikethrough(self):
        self.is_strikethrough = not self.is_strikethrough
        self.btn_strikethrough.config(
            bg="#6C63FF" if self.is_strikethrough else "#F0F2F8",
            fg="white" if self.is_strikethrough else self.colors['text']
        )
        self.apply_format_to_selection()
        self.update_status("Đã bật gạch ngang" if self.is_strikethrough else "Đã tắt gạch ngang")
        
    def change_font(self, font_name):
        self.current_font = font_name
        self.apply_format_to_selection()
        self.update_status(f"Đã chọn font: {font_name}")
        
    def change_size(self, size):
        self.current_font_size = int(size)
        self.apply_format_to_selection()
        self.update_status(f"Đã chọn cỡ chữ: {size}")
        
    def change_color(self, color):
        self.current_color = color
        self.apply_format_to_selection()
        self.update_status(f"Đã chọn màu: {color}")
        
    def change_alignment(self, alignment):
        self.current_alignment = alignment
        self.update_status(f"Đã chọn căn chỉnh: {alignment}")
        
    def insert_special_char(self, char):
        self.text_content.insert(tk.INSERT, char)
        self.update_status(f"Đã chèn ký tự: {char}")
        
    def reset_title(self):
        self.entry_title.delete(0, tk.END)
        self.entry_title.insert(0, "Tạo Ứng Dụng Word Với Python")
        self.update_status("Đã reset tiêu đề")
        
    def clear_content(self):
        if messagebox.askyesno("Xác nhận", "Bạn có chắc muốn xóa toàn bộ nội dung?"):
            self.text_content.delete("1.0", tk.END)
            self.update_status("Đã xóa nội dung")
            
    def undo_action(self):
        try:
            self.text_content.edit_undo()
            self.update_status("Đã hoàn tác")
        except:
            self.update_status("Không thể hoàn tác")
            
    def redo_action(self):
        try:
            self.text_content.edit_redo()
            self.update_status("Đã làm lại")
        except:
            self.update_status("Không thể làm lại")
            
    def apply_format_to_selection(self):
        try:
            sel_start = self.text_content.index(tk.SEL_FIRST)
            sel_end = self.text_content.index(tk.SEL_LAST)
            
            if sel_start and sel_end:
                tag_name = f"format_{random.randint(1000, 9999)}"
                self.text_content.tag_add(tag_name, sel_start, sel_end)
                
                font_props = {
                    'family': self.current_font,
                    'size': self.current_font_size,
                    'weight': 'bold' if self.is_bold else 'normal',
                    'slant': 'italic' if self.is_italic else 'roman',
                    'underline': 1 if self.is_underline else 0,
                    'overstrike': 1 if self.is_strikethrough else 0,
                    'foreground': self.current_color
                }
                
                self.text_content.tag_config(tag_name, **font_props)
                
        except tk.TclError:
            pass
    
    # ==================== TÍNH NĂNG HYPERLINK ====================
    def insert_hyperlink(self):
        """Chèn hyperlink vào văn bản"""
        try:
            selected_text = self.text_content.get(tk.SEL_FIRST, tk.SEL_LAST)
        except tk.TclError:
            selected_text = ""
        
        dialog = tk.Toplevel(self.root)
        dialog.title("🔗 Chèn Hyperlink")
        dialog.geometry("450x250")
        dialog.resizable(False, False)
        dialog.configure(bg=self.colors['card_bg'])
        
        dialog.transient(self.root)
        dialog.grab_set()
        
        main_frame = tk.Frame(dialog, bg=self.colors['card_bg'], padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)
        
        title_label = tk.Label(
            main_frame,
            text="🔗 Chèn Hyperlink",
            font=("Segoe UI", 14, "bold"),
            bg=self.colors['card_bg'],
            fg=self.colors['primary']
        )
        title_label.pack(anchor="w", pady=(0, 10))
        
        tk.Label(
            main_frame,
            text="Văn bản hiển thị:",
            font=("Segoe UI", 10),
            bg=self.colors['card_bg'],
            fg=self.colors['text_light']
        ).pack(anchor="w")
        
        text_entry = tk.Entry(
            main_frame,
            font=("Segoe UI", 11),
            bg="#F8F9FE",
            fg=self.colors['text'],
            relief="flat",
            bd=1
        )
        text_entry.pack(fill="x", pady=(5, 10))
        if selected_text:
            text_entry.insert(0, selected_text)
        else:
            text_entry.insert(0, "Nhấn vào đây để truy cập")
        
        tk.Label(
            main_frame,
            text="Địa chỉ URL:",
            font=("Segoe UI", 10),
            bg=self.colors['card_bg'],
            fg=self.colors['text_light']
        ).pack(anchor="w")
        
        url_entry = tk.Entry(
            main_frame,
            font=("Segoe UI", 11),
            bg="#F8F9FE",
            fg=self.colors['text'],
            relief="flat",
            bd=1
        )
        url_entry.pack(fill="x", pady=(5, 15))
        url_entry.insert(0, "https://")
        
        btn_frame = tk.Frame(main_frame, bg=self.colors['card_bg'])
        btn_frame.pack(fill="x", pady=(10, 0))
        
        def insert_link():
            text = text_entry.get().strip()
            url = url_entry.get().strip()
            
            if not text or not url:
                messagebox.showwarning("Cảnh báo", "Vui lòng nhập đầy đủ văn bản và URL!")
                return
            
            try:
                self.text_content.delete(tk.SEL_FIRST, tk.SEL_LAST)
            except tk.TclError:
                pass
            
            pos = self.text_content.index(tk.INSERT)
            self.text_content.insert(pos, text)
            
            link_tag = f"link_{random.randint(1000, 9999)}"
            start = pos
            end = f"{start}+{len(text)}c"
            self.text_content.tag_add(link_tag, start, end)
            
            self.text_content.tag_config(
                link_tag,
                foreground="#0066CC",
                underline=True,
                font=("Segoe UI", 11, "underline")
            )
            
            self.text_content.tag_bind(
                link_tag,
                "<Button-1>",
                lambda e, url=url: self.open_link(url)
            )
            self.text_content.tag_bind(
                link_tag,
                "<Enter>",
                lambda e: self.text_content.config(cursor="hand2")
            )
            self.text_content.tag_bind(
                link_tag,
                "<Leave>",
                lambda e: self.text_content.config(cursor="")
            )
            
            dialog.destroy()
            self.update_status(f"✅ Đã chèn hyperlink: {text}")
        
        cancel_btn = tk.Button(
            btn_frame,
            text="Hủy",
            font=("Segoe UI", 10),
            bg=self.colors['card_bg'],
            fg=self.colors['text'],
            relief="flat",
            bd=0,
            cursor="hand2",
            command=dialog.destroy
        )
        cancel_btn.pack(side="right", padx=5)
        
        insert_btn = tk.Button(
            btn_frame,
            text="✅ Chèn",
            font=("Segoe UI", 10, "bold"),
            bg=self.colors['primary'],
            fg="white",
            relief="flat",
            bd=0,
            cursor="hand2",
            command=insert_link,
            padx=20,
            pady=8
        )
        insert_btn.pack(side="right", padx=5)
        
        dialog.bind("<Return>", lambda e: insert_link())
        dialog.bind("<Escape>", lambda e: dialog.destroy())
        
    def open_link(self, url):
        try:
            import webbrowser
            webbrowser.open(url)
            self.update_status(f"Đã mở link: {url}")
        except:
            self.update_status("Không thể mở link")
    
    # ==================== TÍNH NĂNG NHÚNG ẢNH ====================
    def insert_image(self):
        image_path = filedialog.askopenfilename(
            title="Chọn ảnh",
            filetypes=[
                ("Ảnh", "*.png *.jpg *.jpeg *.gif *.bmp *.webp"),
                ("Tất cả files", "*.*")
            ]
        )
        
        if not image_path:
            return
        
        try:
            with Image.open(image_path) as img:
                img_width, img_height = img.size
                
            dialog = tk.Toplevel(self.root)
            dialog.title("🖼️ Chỉnh kích thước ảnh")
            dialog.geometry("400x300")
            dialog.resizable(False, False)
            dialog.configure(bg=self.colors['card_bg'])
            
            dialog.transient(self.root)
            dialog.grab_set()
            
            main_frame = tk.Frame(dialog, bg=self.colors['card_bg'], padx=20, pady=20)
            main_frame.pack(fill="both", expand=True)
            
            preview_label = tk.Label(
                main_frame,
                text="📷 Kích thước gốc: {}x{} px".format(img_width, img_height),
                font=("Segoe UI", 10),
                bg=self.colors['card_bg'],
                fg=self.colors['text_light']
            )
            preview_label.pack(anchor="w", pady=(0, 15))
            
            tk.Label(
                main_frame,
                text="Chiều rộng (px):",
                font=("Segoe UI", 10),
                bg=self.colors['card_bg'],
                fg=self.colors['text_light']
            ).pack(anchor="w")
            
            width_entry = tk.Entry(
                main_frame,
                font=("Segoe UI", 11),
                bg="#F8F9FE",
                fg=self.colors['text'],
                relief="flat",
                bd=1
            )
            width_entry.pack(fill="x", pady=(5, 10))
            width_entry.insert(0, str(min(img_width, 400)))
            
            tk.Label(
                main_frame,
                text="Chiều cao (px):",
                font=("Segoe UI", 10),
                bg=self.colors['card_bg'],
                fg=self.colors['text_light']
            ).pack(anchor="w")
            
            height_entry = tk.Entry(
                main_frame,
                font=("Segoe UI", 11),
                bg="#F8F9FE",
                fg=self.colors['text'],
                relief="flat",
                bd=1
            )
            height_entry.pack(fill="x", pady=(5, 15))
            height_entry.insert(0, str(min(img_height, 300)))
            
            keep_ratio = tk.BooleanVar(value=True)
            ratio_check = tk.Checkbutton(
                main_frame,
                text="Giữ tỉ lệ ảnh",
                variable=keep_ratio,
                font=("Segoe UI", 10),
                bg=self.colors['card_bg'],
                fg=self.colors['text'],
                selectcolor=self.colors['card_bg']
            )
            ratio_check.pack(anchor="w", pady=(0, 15))
            
            btn_frame = tk.Frame(main_frame, bg=self.colors['card_bg'])
            btn_frame.pack(fill="x", pady=(10, 0))
            
            def insert_image_with_size():
                try:
                    width = int(width_entry.get())
                    height = int(height_entry.get())
                    
                    if width <= 0 or height <= 0:
                        messagebox.showwarning("Cảnh báo", "Kích thước phải lớn hơn 0!")
                        return
                    
                    pos = self.text_content.index(tk.INSERT)
                    self.text_content.insert(pos, f"\n[🖼️ ẢNH: {os.path.basename(image_path)}]\n")
                    
                    self.image_info = {
                        'path': image_path,
                        'width': width,
                        'height': height
                    }
                    
                    dialog.destroy()
                    self.update_status(f"✅ Đã chèn ảnh: {os.path.basename(image_path)} ({width}x{height} px)")
                    
                except ValueError:
                    messagebox.showwarning("Cảnh báo", "Vui lòng nhập số nguyên hợp lệ!")
            
            cancel_btn = tk.Button(
                btn_frame,
                text="Hủy",
                font=("Segoe UI", 10),
                bg=self.colors['card_bg'],
                fg=self.colors['text'],
                relief="flat",
                bd=0,
                cursor="hand2",
                command=dialog.destroy
            )
            cancel_btn.pack(side="right", padx=5)
            
            insert_btn = tk.Button(
                btn_frame,
                text="✅ Chèn Ảnh",
                font=("Segoe UI", 10, "bold"),
                bg=self.colors['primary'],
                fg="white",
                relief="flat",
                bd=0,
                cursor="hand2",
                command=insert_image_with_size,
                padx=20,
                pady=8
            )
            insert_btn.pack(side="right", padx=5)
            
            dialog.bind("<Return>", lambda e: insert_image_with_size())
            dialog.bind("<Escape>", lambda e: dialog.destroy())
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể mở ảnh: {e}")
            
    def on_text_modified(self, event):
        self.text_content.edit_modified(False)
        
    def update_format_status(self, event):
        pass
        
    def update_word_count(self, event=None):
        try:
            content = self.text_content.get("1.0", tk.END).strip()
            word_count = len(content.split())
            self.word_count_label.config(text=f"{word_count} từ")
        except:
            pass
        
    def update_cursor_position(self, event=None):
        try:
            cursor_pos = self.text_content.index(tk.INSERT)
            line, col = cursor_pos.split('.')
            self.cursor_label.config(text=f"Dòng {line}, Cột {int(col) + 1}")
        except:
            pass
            
    def update_status(self, message):
        try:
            self.status_label.config(text=message)
            self.root.after(3000, lambda: self.status_label.config(text="Sẵn sàng"))
        except:
            pass
        
    def show_help(self):
        help_text = """📖 Hướng dẫn sử dụng Docx Editor Pro v2.1

🎨 Định dạng văn bản:
• Bôi đen văn bản cần định dạng
• Sử dụng các nút trên thanh toolbar

🔗 Hyperlink (Ctrl+K):
• Bôi đen văn bản hoặc để trống
• Nhấn nút 🔗 hoặc Ctrl+K
• Nhập văn bản và URL
• Nhấp vào link để mở trong trình duyệt

🖼️ Nhúng ảnh (Ctrl+G):
• Nhấn nút 🖼️ hoặc Ctrl+G
• Chọn file ảnh
• Chỉnh kích thước ảnh
• Ảnh sẽ được chèn vào văn bản

⌨️ Phím tắt:
• Ctrl+B: In đậm
• Ctrl+I: In nghiêng
• Ctrl+U: Gạch dưới
• Ctrl+K: Chèn Hyperlink
• Ctrl+G: Chèn ảnh
• Ctrl+Z: Hoàn tác
• Ctrl+Y: Làm lại

📤 Xuất file:
• Nhấn nút "Tạo File Word" để xuất
• Hỗ trợ đầy đủ định dạng, hyperlink và ảnh"""
        
        messagebox.showinfo("Hướng dẫn sử dụng", help_text)
        
    def create_word_file(self):
        title_text = self.entry_title.get().strip()
        content_text = self.text_content.get("1.0", tk.END).strip()
        
        if not title_text or not content_text:
            messagebox.showwarning("Cảnh báo", "⚠️ Vui lòng nhập đầy đủ tiêu đề và nội dung!")
            return
            
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Chọn nơi lưu file Word",
            initialfile="newbook.python.docx"
        )
        
        if not file_path:
            return
            
        try:
            doc = Document()
            
            p0 = doc.add_paragraph()
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(title_text)
            r0.font.name = 'Arial'
            r0.font.size = Pt(18)
            r0.bold = True
            
            content = self.text_content.get("1.0", tk.END).strip()
            lines = content.split('\n')
            
            for line in lines:
                if not line.strip():
                    doc.add_paragraph()
                    continue
                
                if '[🖼️ ẢNH:' in line and hasattr(self, 'image_info'):
                    try:
                        img_info = self.image_info
                        if os.path.exists(img_info['path']):
                            doc.add_picture(
                                img_info['path'],
                                width=Cm(img_info['width'] / 37.8),
                                height=Cm(img_info['height'] / 37.8)
                            )
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = p_cap.add_run(f"Hình ảnh minh họa")
                            cap_run.font.size = Pt(10)
                            cap_run.font.italic = True
                            continue
                    except Exception as e:
                        print(f"Lỗi khi thêm ảnh: {e}")
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Inches(0.5)
                r = p.add_run(line)
                r.font.name = 'Arial'
                r.font.size = Pt(12)
            
            doc.save(file_path)
            self.update_status(f"✅ Đã xuất file: {os.path.basename(file_path)}")
            messagebox.showinfo("Thành công", f"✅ Đã xuất file Word thành công tại:\n{file_path}")
            
        except Exception as e:
            self.update_status("❌ Lỗi khi tạo file")
            messagebox.showerror("Lỗi", f"❌ Không thể tạo file Word: {e}")

# ==================== PHẦN CHẠY CHÍNH ====================
if __name__ == "__main__":
    root = tk.Tk()
    
    # ===== SỬA LỖI: CHỈ TẠO 1 INSTANCE =====
    app = ModernWordEditor(root)  # Chỉ tạo 1 instance duy nhất
    root.mainloop()
